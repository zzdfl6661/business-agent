"""
文档加载与切分
==============
- load_documents : 按扩展名加载（md/txt → TextLoader；pdf → PyPDFLoader）
- split_documents: RecursiveCharacterTextSplitter（中文按段落/句子切，chunk 500 / overlap 80）
- ingest         : 加载 → 切分 → 向量化入库（幂等：先清同名 doc_type）

示例知识文档位于 rag/data/（门店运营SOP.md / 推广优化策略.md / 活动运营规则.md / ROI异常分析案例.md），
服务启动时自动 ingest。
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.settings import PROJECT_DIR
from rag.retriever import get_vector_client

logger = logging.getLogger(__name__)

DATA_DIR = PROJECT_DIR / "rag" / "data"

# 文档名 → doc_type（用于元数据与过滤）
DOC_TYPE_MAP = {
    "门店运营SOP": "sop",
    "推广优化策略": "strategy",
    "活动运营规则": "activity",
    "ROI异常分析案例": "case",
    "满意度回访话术": "callback",
    "门店晋升制度": "hr",
    "门店薪资绩效管理办法": "salary",
    "员工手册": "handbook",
    "公司背景": "company",
    "门店日常工作安排": "daily",
    "公司高管": "org",
}


def _doc_type_for(path: Path) -> str:
    for keyword, doc_type in DOC_TYPE_MAP.items():
        if keyword in path.stem:
            return doc_type
    return "general"


def load_documents(paths: list[Path]) -> list[Document]:
    """按扩展名加载文档。md/txt 用 TextLoader；pdf 用 PyPDFLoader（可选依赖）。"""
    docs: list[Document] = []
    for path in paths:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            try:
                from langchain_community.document_loaders import PyPDFLoader

                loaded_pdf = PyPDFLoader(str(path)).load()
                for doc in loaded_pdf:
                    doc.metadata["doc_type"] = _doc_type_for(path)
                    doc.metadata["source"] = path.name
                docs.extend(loaded_pdf)
                logger.info("已加载 PDF: %s（%s）", path.name, _doc_type_for(path))
                continue
            except ImportError:
                logger.warning("未安装 pypdf，跳过 %s", path.name)
                continue

        if suffix == ".docx":
            try:
                from docx import Document as DocxDocument

                d = DocxDocument(str(path))
                parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
                # 提取表格（制度类文档关键信息常在表格中）
                for tbl in d.tables:
                    for row in tbl.rows:
                        cells = [c.text.strip().replace("\n", " ") for c in row.cells if c.text.strip()]
                        if cells:
                            parts.append(" | ".join(cells))
                text = "\n".join(parts)
                if text.strip():
                    doc = Document(page_content=text, metadata={
                        "doc_type": _doc_type_for(path), "source": path.name,
                    })
                    docs.append(doc)
                    logger.info("已加载 DOCX: %s（%s，%s 字符）", path.name, _doc_type_for(path), len(text))
                continue
            except ImportError:
                logger.warning("未安装 python-docx，跳过 %s", path.name)
                continue

        from langchain_community.document_loaders import TextLoader

        loaded = TextLoader(str(path), encoding="utf-8").load()
        for doc in loaded:
            doc.metadata["doc_type"] = _doc_type_for(path)
            doc.metadata["source"] = path.name
        docs.extend(loaded)
        logger.info("已加载文档: %s（%s）", path.name, _doc_type_for(path))
    return docs


def split_documents(docs: list[Document], chunk_size: int = 400, chunk_overlap: int = 80) -> list[Document]:
    """
    文本切分（滑动窗口）：
    - chunk_size=400：中文 1 字 ≈ 1 token，all-MiniLM-L6-v2 输入上限 256 tokens，
      500 字会被截断，400 字可完整嵌入且上下文足够
    - chunk_overlap=80：滑动窗口重叠，保证跨块语义（句子）不丢失
    - separators 中文优先按段落/句子切
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", " "],  # 中文优先按段落/句子
    )
    return splitter.split_documents(docs)


# 父子切割参数：子块（检索精准，对齐 embedding 上限）→ 父块（上下文完整供 LLM）
CHILD_CHUNK_SIZE = 350
CHILD_CHUNK_OVERLAP = 50
PARENT_CHUNK_SIZE = 1200
PARENT_CHUNK_OVERLAP = 150
PARENT_MAX_CHARS = 2500  # 父块上限：超过则按次级标题/字数细分（防超大章节撑爆 token）

# 章节标题识别：第X章/篇/节、中文序号"一、"（不匹配阿拉伯数字"1、"——正文列表序号会误伤）
_HEADER_RE = re.compile(r"^\s*((第[一二三四五六七八九十百]+[章篇节])|([一二三四五六七八九十]+、))\s*\S")


def _split_parents_by_headers(text: str, max_chars: int = PARENT_MAX_CHARS) -> list[str]:
    """
    结构化父块切分：父块 = 章节板块（用户诉求：检索只返回到对应板块，不返回整份大文档）。
    1. 按章节标题（第X章 / 一、 / 1.2）切出板块
    2. 板块 > max_chars 时：内部按次级标题切，仍超长则按字符硬切
    """
    blocks: list[tuple[str, list[str]]] = []
    current_title = "文档开头"
    current_lines: list[str] = []
    for line in text.split("\n"):
        line_s = line.strip()
        is_header = bool(_HEADER_RE.match(line_s)) and len(line_s) < 40
        if is_header:
            if current_lines:
                blocks.append((current_title, current_lines))
            current_title = line_s
            current_lines = [line_s]
        else:
            current_lines.append(line_s)
    if current_lines:
        blocks.append((current_title, current_lines))

    parents: list[str] = []
    for title, lines in blocks:
        block = "\n".join(l for l in lines if l).strip()
        if not block:
            continue
        if len(block) <= max_chars:
            parents.append(block)
            continue
        # 超长板块：按次级标题（X.X / 一、）再切
        sub = _split_sub_headers(block, max_chars)
        parents.extend(sub)
    return parents


def _split_sub_headers(text: str, max_chars: int) -> list[str]:
    """板块内按次级标题切分，仍超长则按 max_chars 硬切。"""
    out: list[str] = []
    cur: list[str] = []
    for line in text.split("\n"):
        line_s = line.strip()
        is_sub = bool(re.match(r"^\s*\d+(\.\d+)+\s", line_s)) and len(line_s) < 60
        if is_sub and cur:
            out.append("\n".join(cur).strip())
            cur = [line_s]
        else:
            cur.append(line_s)
    if cur:
        out.append("\n".join(cur).strip())

    final: list[str] = []
    for block in out:
        if len(block) <= max_chars:
            final.append(block)
        else:
            # 硬切（滑动窗口）
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=max_chars, chunk_overlap=150, separators=["\n\n", "\n", "。", "！", "？", " "]
            )
            final.extend(s for s in splitter.split_text(block) if s.strip())
    return final


def split_documents_hierarchical(
    docs: list[Document],
    child_size: int = CHILD_CHUNK_SIZE,
    child_overlap: int = CHILD_CHUNK_OVERLAP,
    parent_size: int = PARENT_CHUNK_SIZE,
    parent_overlap: int = PARENT_CHUNK_OVERLAP,
) -> list[Document]:
    """
    父子切割（Parent-Child Chunking）：
    - 父块：按章节标题结构化切分（第X章/一、/1.2 为板块边界），上下文完整供 LLM 引用；
      检索命中子块只返回对应板块，不返回整份大文档
    - 子块（350 字）：父块内再切，语义聚焦，供 embedding/BM25 精准检索
    - 子块 metadata 携带 parent_id + parent_content；检索命中子块后由 retriever 回取父块
    """
    parent_splitter = RecursiveCharacterTextSplitter(
        chunk_size=parent_size, chunk_overlap=parent_overlap,
        separators=["\n\n", "\n", "。", "！", "？", " "],
    )
    child_splitter = RecursiveCharacterTextSplitter(
        chunk_size=child_size, chunk_overlap=child_overlap,
        separators=["\n\n", "\n", "。", "！", "？", " "],
    )
    chunks: list[Document] = []
    for doc in docs:
        parents = _split_parents_by_headers(doc.page_content)
        if not parents:  # 空文档兜底：纯字数切
            parents = parent_splitter.split_text(doc.page_content)
        base_meta = dict(doc.metadata)
        for pi, ptext in enumerate(parents):
            pid = f"{base_meta.get('source', 'doc')}#p{pi}"
            children = child_splitter.split_text(ptext)
            for ci, ctext in enumerate(children):
                chunks.append(Document(
                    page_content=ctext,
                    metadata={
                        **base_meta,
                        "parent_id": pid,
                        "parent_content": ptext,
                        "child_index": ci,
                    },
                ))
    return chunks


def get_sample_docs() -> list[Path]:
    """rag/data/ 下的示例知识文档（md/txt/pdf/docx）。"""
    return sorted(p for p in DATA_DIR.glob("*.*") if p.is_file() and p.suffix.lower() in (".md", ".txt", ".pdf", ".docx"))


# ---------------- 上传入库（/api/rag/upload 专用） ----------------
# 上传安全约束（修复路径穿越 + 幂等粒度错误，#2）：
# - 文件名取 basename（拒绝 ../../x 等穿越路径）
# - 白名单后缀：md/txt/pdf/docx
# - 大小上限 20MB（防打满磁盘）
# - 内容非空校验
# - 幂等删除按 source（文件名）而非 doc_type——上传 general 文档不再清空其他 general 文档
ALLOWED_UPLOAD_EXTS = (".md", ".txt", ".pdf", ".docx")
MAX_UPLOAD_BYTES = 20 * 1024 * 1024  # 20MB


def upload_and_ingest(filename: str, content: bytes) -> dict:
    """上传知识文档：安全校验 → 落盘 → 切分 → 向量化入库。

    安全设计：
    - 文件名取 basename（防路径穿越），扩展名白名单校验；
    - 内容非空校验、大小 ≤ MAX_UPLOAD_BYTES（20MB）；
    - 幂等：只清理本次文件的旧数据（按 source=文件名），
      **不再按 doc_type 清库**——修复"上传 general 文档清空所有 general 文档"的误伤。

    返回 {success, file, chunks, doc_type}；校验/入库失败抛 ValueError/异常，由调用方返回错误。
    """
    name = Path(filename or "").name  # basename：防路径穿越（"../../x.md" → "x.md"）
    suffix = Path(name).suffix.lower()
    if not name or name in (".", ".."):
        raise ValueError("文件名无效")
    if suffix not in ALLOWED_UPLOAD_EXTS:
        raise ValueError(
            f"不支持的文件类型：{name or '(空)'}"
            f"（仅支持 {'/'.join(sorted(ALLOWED_UPLOAD_EXTS))}）"
        )
    if not content or not content.strip():
        raise ValueError("文件内容为空，拒绝入库")
    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError(f"文件超过大小限制（{MAX_UPLOAD_BYTES // 1024 // 1024}MB）")

    dest = DATA_DIR / name
    dest.write_bytes(content)
    docs = load_documents([dest])
    if not docs:
        raise ValueError("文档解析后无有效内容，拒绝入库")
    chunks = split_documents_hierarchical(docs)
    if not chunks:
        raise ValueError("文档切分后无有效 chunk，拒绝入库")

    client = get_vector_client()
    # 幂等：按 source（文件名）精确清理旧数据，不影响其他文档
    client.delete(None, source=name)
    client.add_documents(chunks, None)

    doc_types = sorted({d.metadata.get("doc_type") for d in docs} - {None})
    logger.info("上传入库完成：%s（%s chunks，doc_type=%s）", name, len(chunks), doc_types)
    return {"success": True, "file": name, "chunks": len(chunks), "doc_type": doc_types}


def ingest(paths: list[Path] | None = None, doc_type: str | None = None) -> int:
    """加载 → 切分 → 入库。返回入库 chunk 数；失败时抛出（调用方决定降级）。

    幂等：只删除本次文档涉及的 doc_type（不清全量，避免误删 report 经验层）。
    """
    paths = paths or get_sample_docs()
    docs = load_documents(paths)
    if not docs:
        return 0
    chunks = split_documents_hierarchical(docs)
    client = get_vector_client()
    # 幂等：按本次涉及的 doc_type 清理（doc_type 参数优先，否则按文档 metadata）
    target_types = {doc_type} if doc_type else {d.metadata.get("doc_type") for d in docs if d.metadata.get("doc_type")}
    for dt in target_types:
        client.delete(dt)
    client.add_documents(chunks, None)
    logger.info("RAG ingest 完成：%s 个 chunk（父子切割，来源 %s 个文件）", len(chunks), len(paths))
    return len(chunks)
